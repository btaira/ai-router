"""Extracts plain text from an uploaded document so it can be prepended to
a prompt. Deliberately text-only (no vision/multimodal support) — every
provider adapter takes a single text prompt, so a document becomes context
the model reads, not an image it sees.
"""
from __future__ import annotations

import io
from pathlib import Path

MAX_EXTRACTED_CHARS = 60_000  # ~15k tokens; keeps six parallel stage-1 calls affordable

_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log", ".yaml", ".yml"}


class UnsupportedFileType(Exception):
    pass


class ExtractionError(Exception):
    pass


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001 - surfaced to the user as a plain extraction failure
        raise ExtractionError(f"could not read PDF: {exc}") from exc
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one bad page shouldn't kill the whole extraction
            continue
    text = "\n\n".join(pages).strip()
    if not text:
        raise ExtractionError("no extractable text found in PDF (it may be scanned/image-only)")
    return text


def _extract_docx(content: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"could not read .docx: {exc}") from exc
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("no extractable text found in .docx")
    return text


def _extract_plain_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="replace")


def extract_text(filename: str, content: bytes) -> tuple[str, bool]:
    """Returns (text, truncated). Raises UnsupportedFileType / ExtractionError."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        text = _extract_pdf(content)
    elif ext == ".docx":
        text = _extract_docx(content)
    elif ext in _TEXT_EXTENSIONS or ext == "":
        text = _extract_plain_text(content)
    else:
        raise UnsupportedFileType(
            f"unsupported file type {ext!r} — supported: .txt, .md, .csv, .json, .yaml, .log, .pdf, .docx"
        )

    truncated = len(text) > MAX_EXTRACTED_CHARS
    if truncated:
        text = text[:MAX_EXTRACTED_CHARS]
    return text, truncated
